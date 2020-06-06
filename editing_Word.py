import docx

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
print(p.runs, '\n')

p.runs[1].bold = True
d.save('demo5.docx')

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))