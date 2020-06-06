import docx

d = docx.Document('demo.docx')
print(d.paragraphs, '\n')
print(d.paragraphs[0].text, '\n')

p = d.paragraphs[1]
print(p.runs, '\n')
print(p.runs[0].text, '\n')

for i in range(len(p.runs)):
    print(p.runs[i].text, '\n')

print(p.runs[1].bold, '\n')
print(p.runs[3].italic, '\n')

p.runs[3].underline = True
p.runs[3].text = 'italic and underline'
d.save('demo2.docx')

p.style = 'Title'
d.save('demo3.docx')